"""
docx_parser — python-docx 抽取全文（含表格转 markdown 管道格式，保留列语义）

表格必须转 markdown 而非纯文本拼接，否则 AI 丢失"参数名 | 类型 | 必填 | 说明"的列对应关系。
"""

import docx


def _paragraphs_to_text(doc) -> str:
    lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(lines)


def _tables_to_markdown(doc) -> str:
    blocks: list[str] = []
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        header_cells = [c.text.strip() for c in rows[0].cells]
        blocks.append("| " + " | ".join(header_cells) + " |")
        blocks.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            blocks.append("| " + " | ".join(cells) + " |")
        blocks.append("")
    return "\n".join(blocks)


def extract_text_docx(path: str) -> str:
    """抽取 docx 全文（段落 + 表格 markdown）。失败抛异常由调用方处理。"""
    doc = docx.Document(path)
    parts = [_paragraphs_to_text(doc), _tables_to_markdown(doc)]
    return "\n\n".join(p for p in parts if p).strip()
